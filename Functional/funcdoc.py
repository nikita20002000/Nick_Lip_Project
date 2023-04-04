from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
import datetime
import pictures


def create_titul(name="No_name", surname="No_surname", lastname="No_lastname", group="No_group", position_in_sub_list=0):
    mass_of_subject_data = [
        {"TUTOR": "Шевцова Инесса Витальевна", "GRADE": "к.т.н., доцент", "CAFEDRA": "Кафедра математических методов и информационных технологий в управлении", "SUBJECT": "Большие данные и искусственный интеллект в управлении"},
        {"TUTOR": "Дудихин Виктор Владимирович", "GRADE": "к.т.н., доцент", "CAFEDRA": "Кафедра математических методов и информационных технологий в управлении", "SUBJECT": "Введение в искусственный интеллект"},
        {"TUTOR": "Наумов Александр Олегович", "GRADE": "доктор исторических наук", "CAFEDRA": "Кафедра международных организаций и проблем глобального управления", "SUBJECT": "Государственное управление информационной политикой в цифровую эпоху"},
        {"TUTOR": "Депелян Рузанна Амбарцумовна", "GRADE": "старший преподаватель", "CAFEDRA": "Кафедра иностранных языков факультета государственного управления", "SUBJECT": "Иностранный язык для делового общения"},
        {"TUTOR": "Кудина Марианна Валерьевна", "GRADE": "Доктор экономических наук", "CAFEDRA": "Кафедра менеджмента", "SUBJECT": "Инновационная экономика"},
        {"TUTOR": "Шевцова Инесса Витальевна", "GRADE": "к.т.н., доцент", "CAFEDRA": "Кафедра математических методов и информационных технологий в управлении", "SUBJECT": "Информационно-аналитические технологии в государственном и муниципальном управлении"},
        {"TUTOR": "Зайцева Татьяна Вячеславовна", "GRADE": "Доктор экономических наук", "CAFEDRA": "Кафедра управления персоналом", "SUBJECT": "Кадровая политика и кадровый аудит"},
        {"TUTOR": "Наумов Александр Олегович", "GRADE": "доктор исторических наук", "CAFEDRA": "Кафедра международных организаций и проблем глобального управления", "SUBJECT": "Лидерство"},
        {"TUTOR": "Глазьев Сергей Юрьевич", "GRADE": "Доктор экономических наук, профессор, академик РАН", "CAFEDRA": "Кафедра экономической политики и экономических измерений", "SUBJECT": "Макроэкономическая политика"},
        {"TUTOR": "Никитин Александр Иванович", "GRADE": "доктор политических наук", "CAFEDRA": "Кафедра международных организаций и проблем глобального управления", "SUBJECT": "Международная безопасность"},
        {"TUTOR": "Фамилия Имя Отчество", "GRADE": "ученая степень", "CAFEDRA": "Кафедра", "SUBJECT": "предмет"},
        {"TUTOR": "Фамилия Имя Отчество", "GRADE": "ученая степень", "CAFEDRA": "Кафедра", "SUBJECT": "предмет"},
        {"TUTOR": "Лексин Иван Владимирович", "GRADE": "профессор", "CAFEDRA": "Кафедра правовых основ управления", "SUBJECT": "Правовое обеспечение государственного и муниципального управления"},
        {"TUTOR": "Попова Светлана Сергеевна", "GRADE": "к.ю.н., доцент", "CAFEDRA": "Кафедра правовых основ управления", "SUBJECT": "Правовые аспекты использования искусственного интеллекта в государственном и муниципальном управлении"},
        {"TUTOR": "Фамилия Имя Отчество", "GRADE": "ученая степень", "CAFEDRA": "Кафедра", "SUBJECT": "предмет"},
        {"TUTOR": "Ведута Елена Николаевна", "GRADE": "доктор экономических наук", "CAFEDRA": "Кафедра стратегического планирования и экономической политики", "SUBJECT": "Стратегическое планирование"},
        {"TUTOR": "Купряшин Геннадий Львович", "GRADE": "доктор политических наук", "CAFEDRA": "Кафедра теории и методологии государственного и муниципального управления", "SUBJECT": "Теория и механизмы современного государственного управления"},
        {"TUTOR": "Терентьева Ольга Игоревна", "GRADE": "к.э.н., доцент", "CAFEDRA": "Кафедра экономики инновационного развития", "SUBJECT": "Управление цифровой экономикой"},
        {"TUTOR": "Купряшин Геннадий Львович", "GRADE": "доктор политических наук", "CAFEDRA": "Кафедра теории и методологии государственного и муниципального управления", "SUBJECT": "Управление государственными программами"},
        {"TUTOR": "Тышкевич Виктория Петровна", "GRADE": "к.э.н., доцент", "CAFEDRA": "Кафедра Экономической и финансовой стратегии", "SUBJECT": "Философия"},
        {"TUTOR": "Опарина Ольга Игоревна", "GRADE": "старший преподаватель", "CAFEDRA": "Кафедра английского языка", "SUBJECT": "Цифровые технологии и искусственный интеллект в управлении персоналом государственных и коммерческих организаций"},
        {"TUTOR": "Каширова Анна Владимировна", "GRADE": "к.э.н., доцент", "CAFEDRA": "Кафедра экономики инновационного развития", "SUBJECT": "Экономика общественного сектора"}
    ]

    doc = Document()
    style = doc.styles['Normal']

    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.font.bold = True
    WD_LINE_SPACING.EXACTLY = 1.5
    with open("MSU_LOGO.jpg", "rb") as img:
        p = doc.add_picture(img)

    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph('МОСКОВСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ', style)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.bold = True
    p.paragraph_format.line_spacing = WD_LINE_SPACING.EXACTLY

    p = doc.add_paragraph('имени М.В.ЛОМОНОСОВА', style)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.bold = True
    p.paragraph_format.line_spacing = WD_LINE_SPACING.EXACTLY

    p = doc.add_paragraph("ФАКУЛЬТЕТ ГОСУДАРСТВЕННОГО УПРАВЛЕНИЯ", style)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.bold = True
    p.paragraph_format.line_spacing = WD_LINE_SPACING.EXACTLY



    p = doc.add_paragraph(f"{mass_of_subject_data[position_in_sub_list]['CAFEDRA']}", style)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.bold = True
    p.paragraph_format.line_spacing = WD_LINE_SPACING.EXACTLY

    p = doc.add_paragraph(f"{mass_of_subject_data[position_in_sub_list]['SUBJECT']}", style)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.bold = True
    p.paragraph_format.line_spacing = WD_LINE_SPACING.EXACTLY

    k = doc.add_paragraph()
    k.paragraph_format.line_spacing = WD_LINE_SPACING.EXACTLY
    k.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    k = doc.add_paragraph()
    k.paragraph_format.line_spacing = WD_LINE_SPACING.EXACTLY
    k.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    k = doc.add_paragraph()
    k.paragraph_format.line_spacing = WD_LINE_SPACING.EXACTLY
    k.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    k = doc.add_paragraph("Выполнил:", style)
    k.bold = True
    k.paragraph_format.line_spacing = WD_LINE_SPACING.EXACTLY
    k.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    p = doc.add_paragraph(f"Студент группы {group}", style)
    p.paragraph_format.line_spacing = WD_LINE_SPACING.EXACTLY
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    p = doc.add_paragraph(f"{surname} {name} {lastname}", style)
    p.paragraph_format.line_spacing = WD_LINE_SPACING.EXACTLY
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    p = doc.add_paragraph("Проверила:", style)
    p.paragraph_format.line_spacing = WD_LINE_SPACING.EXACTLY
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    p = doc.add_paragraph(f"{mass_of_subject_data[position_in_sub_list]['GRADE']}", style)
    p.paragraph_format.line_spacing = WD_LINE_SPACING.EXACTLY
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    p = doc.add_paragraph(f"{mass_of_subject_data[position_in_sub_list]['TUTOR']}", style)
    p.paragraph_format.line_spacing = WD_LINE_SPACING.EXACTLY
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    year = datetime.datetime.strftime(datetime.datetime.now(), '%Y')            #Добавил год в титульник
    p = doc.add_paragraph(f"{year}", style)
    p.paragraph_format.line_spacing = WD_LINE_SPACING.EXACTLY
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save('your_doc.docx')
    return 1


